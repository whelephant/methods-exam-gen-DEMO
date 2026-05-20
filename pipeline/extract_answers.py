"""Examiner-report answer extraction.

Examiner reports come in two forms:
  - .docx (2020+) — VCAA's modern format. Math is stored as MathType OLE objects
                    (rendered to WMF previews inside the docx), which python-docx
                    and pandoc both drop in plain-text extraction.
  - .pdf  (2016-2019) — older scanned reports. (Path not yet implemented.)

docx path:
  1. Convert .docx → .pdf via headless LibreOffice (cached, $0 API).
  2. Locate every "Question Xa." heading in the rendered PDF.
  3. For each heading, crop the page region between this heading and the next as
     a 200 DPI PNG and save under assets/answers/<question_id>.png. Multi-page
     sections are stitched vertically using PIL.
  4. Walk the docx via python-docx to extract the body-text commentary for each
     question section (stored as commentary_md, vector text).
  5. Upsert into the answers table.

The math/worked-answer images become the answer artifact; the body text becomes
the verbatim commentary. Combined, this matches the project rule of "examiner
commentary only" (no LLM-generated solutions).

CLI:
  python -m pipeline.extract_answers --year 2023 --paper 1
  python -m pipeline.extract_answers --year 2023 --paper 1 --force
"""
from __future__ import annotations

import argparse
import io
import json
import re
import subprocess
import sys
from pathlib import Path
from typing import Optional

import fitz
from PIL import Image
from docx import Document

from pipeline.db import REPO_ROOT, connect

ANSWERS_DIR = REPO_ROOT / "assets" / "answers"
ANSWERS_DIR.mkdir(parents=True, exist_ok=True)
MC_ROW_THUMBS_DIR = REPO_ROOT / "assets" / "mc_row_thumbs"
MC_ROW_THUMBS_DIR.mkdir(parents=True, exist_ok=True)

CONVERT_CACHE = REPO_ROOT / "pipeline" / "cache" / "answer_pdfs"
CONVERT_CACHE.mkdir(parents=True, exist_ok=True)

SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

ANSWER_DPI = 200

# Regex matches "Question N", "Question Na.", "Question N.a.i", "Question Nci." (no dot), etc.
# 2023 P1 docx uses "Question 1a." (Heading 4 style). 2023 P2 docx uses "Question 1ci." (no dot
# between letter and roman, often under Heading 3 style or "Normal"). We accept both forms; the
# `_parse_heading` normaliser converts "ci" → "c.i" so it matches the question-extraction format.
QUESTION_HEADING_RE = re.compile(
    r"^Question\s+(\d+)\s*([a-z](?:\.?[ivxlc]+)?)?\s*\.?\s*$",
    re.IGNORECASE,
)


# ─── docx → pdf via LibreOffice ───────────────────────────────────────

def convert_docx_to_pdf(docx_path: Path, *, force: bool = False) -> Path:
    """Headless LibreOffice convert. Result cached under pipeline/cache/answer_pdfs/."""
    out_pdf = CONVERT_CACHE / (docx_path.stem + ".pdf")
    if out_pdf.exists() and not force:
        return out_pdf
    subprocess.run(
        [SOFFICE, "--headless", "--convert-to", "pdf",
         "--outdir", str(CONVERT_CACHE), str(docx_path)],
        check=True, capture_output=True, text=True,
    )
    if not out_pdf.exists():
        raise RuntimeError(f"LibreOffice did not produce {out_pdf}")
    return out_pdf


# ─── locate question headings in the PDF ──────────────────────────────

def _parse_heading(text: str) -> Optional[tuple[int, Optional[str]]]:
    """Returns (question_number, part) or None.

    Normalises the part label so "ci" → "c.i" (2023 P2 docx convention) matches the
    question-extraction format used elsewhere in the DB. The question regex captures
    one letter optionally followed by `?[ivxlc]+`; we re-insert the dot here when missing.
    """
    t = text.strip()
    if not t:
        return None
    m = QUESTION_HEADING_RE.match(t)
    if not m:
        return None
    qn = int(m.group(1))
    part = m.group(2)
    if part:
        part = part.lower()
        # If part is "ci", "cii", "ciii", etc. — a single letter followed by Roman digits
        # with no dot in between — insert the dot so it matches the question_extraction format.
        if len(part) > 1 and "." not in part and part[1] in "ivxlc":
            part = part[0] + "." + part[1:]
    return qn, part


def find_headings(pdf_path: Path) -> list[dict]:
    """Returns list of dicts with question_number, part, page (1-based), y0, y1."""
    doc = fitz.open(pdf_path)
    headings: list[dict] = []
    try:
        for pn in range(doc.page_count):
            pg = doc.load_page(pn)
            for x0, y0, x1, y1, text, *_ in pg.get_text("blocks"):
                t = text.strip()
                # Headings start with "Question " and are short.
                if not t.startswith("Question") or len(t) > 30:
                    continue
                parsed = _parse_heading(t)
                if parsed is None:
                    continue
                qn, part = parsed
                headings.append({
                    "question_number": qn,
                    "part": part,
                    "page": pn + 1,
                    "y0": float(y0),
                    "y1": float(y1),
                    "text": t,
                })
    finally:
        doc.close()
    return headings


# ─── crop region (possibly multi-page) into one PNG ───────────────────

# Top/bottom padding to avoid running headers/footers when section spans pages.
PAGE_HEADER_SKIP = 50
PAGE_FOOTER_SKIP = 50


def crop_section(pdf_path: Path, ranges: list[tuple[int, float, float]],
                 out_path: Path, dpi: int = ANSWER_DPI) -> None:
    """Render the union of (page, y0, y1) ranges into one stitched PNG."""
    doc = fitz.open(pdf_path)
    try:
        tiles: list[Image.Image] = []
        for (page_num, y0, y1) in ranges:
            pg = doc.load_page(page_num - 1)
            if y1 - y0 < 5:
                continue
            rect = fitz.Rect(0, y0, pg.rect.width, y1)
            pix = pg.get_pixmap(clip=rect, dpi=dpi)
            tiles.append(Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB"))
        if not tiles:
            raise RuntimeError(f"no tiles produced for {out_path}")
        if len(tiles) == 1:
            tiles[0].save(out_path, optimize=True)
            return
        # Stitch vertically; align to widest tile's width.
        max_w = max(t.width for t in tiles)
        total_h = sum(t.height for t in tiles)
        canvas = Image.new("RGB", (max_w, total_h), "white")
        y = 0
        for t in tiles:
            canvas.paste(t, (0, y))
            y += t.height
        canvas.save(out_path, optimize=True)
    finally:
        doc.close()


def section_ranges(pdf_path: Path, headings: list[dict]) -> list[tuple[dict, list[tuple[int, float, float]]]]:
    """For each heading, compute its page-range list (page, y0, y1) until the next heading."""
    doc = fitz.open(pdf_path)
    try:
        page_heights = {pn + 1: doc.load_page(pn).rect.height for pn in range(doc.page_count)}
    finally:
        doc.close()
    out: list[tuple[dict, list[tuple[int, float, float]]]] = []
    for i, h in enumerate(headings):
        start_page = h["page"]
        start_y = h["y1"]
        if i + 1 < len(headings):
            end_page = headings[i + 1]["page"]
            end_y = headings[i + 1]["y0"]
        else:
            end_page = max(page_heights.keys())
            end_y = page_heights[end_page] - PAGE_FOOTER_SKIP
        if end_page == start_page:
            ranges = [(start_page, start_y, end_y)]
        else:
            ranges = [(start_page, start_y, page_heights[start_page] - PAGE_FOOTER_SKIP)]
            for p in range(start_page + 1, end_page):
                ranges.append((p, PAGE_HEADER_SKIP, page_heights[p] - PAGE_FOOTER_SKIP))
            ranges.append((end_page, PAGE_HEADER_SKIP, end_y))
        out.append((h, ranges))
    return out


# ─── commentary text via python-docx ──────────────────────────────────

def pdf_commentary_by_question(pdf_path: Path) -> dict[tuple[int, Optional[str]], str]:
    """PDF-native counterpart to `docx_commentary_by_question`.

    For 2016–2019 reports the source is already a PDF (VCAA didn't publish docx for
    those years). We can't use python-docx, so we segment by the same 'Question Xa.'
    heading scheme but pull prose via PyMuPDF text extraction.

    Math content extracts as gibberish (CID-encoded glyphs), same as on the source
    papers — that's accepted because the cropped image is the canonical answer
    artefact; commentary_md is just for search/audit.
    """
    import fitz as _fitz
    PROSE_MIN_CHARS = 25
    src = _fitz.open(pdf_path)
    try:
        headings = find_headings(pdf_path)
        sections: dict[tuple[int, Optional[str]], list[str]] = {}
        for i, h in enumerate(headings):
            key = (h["question_number"], h["part"])
            sections.setdefault(key, [])
            start_page = h["page"]  # 1-based from find_headings
            start_y = h["y1"]  # below the heading line
            if i + 1 < len(headings):
                nxt = headings[i + 1]
                end_page, end_y = nxt["page"], nxt["y0"]
            else:
                end_page = src.page_count
                end_y = src.load_page(end_page - 1).rect.height
            for p in range(start_page, end_page + 1):
                page = src.load_page(p - 1)  # 0-based for fitz
                clip_y0 = start_y if p == start_page else 0
                clip_y1 = end_y if p == end_page else page.rect.height
                clip = _fitz.Rect(0, clip_y0, page.rect.width, clip_y1)
                # 'blocks' groups paragraph-level text including CID-jumbled math
                # rows; we filter math fragments by length to keep prose-like blocks.
                for block in page.get_text("blocks", clip=clip):
                    text = (block[4] or "").strip()
                    if not text or len(text) < PROSE_MIN_CHARS:
                        continue
                    # Skip running headers / footers / heading repeats.
                    low = text.lower()
                    if "examination report" in low or "© vcaa" in low or low.startswith("page "):
                        continue
                    if _parse_heading(text) is not None:
                        continue
                    sections[key].append(text)
        return {k: "\n\n".join(v).strip() for k, v in sections.items()}
    finally:
        src.close()


def docx_commentary_by_question(docx_path: Path) -> dict[tuple[int, Optional[str]], str]:
    """Walk paragraphs, segment by 'Question Xa.' headings, return body text per key.

    The VCAA report author inconsistently applies paragraph styles — prose paragraphs
    are sometimes under "VCAA body", sometimes under "VCAA math equations" (a style
    that's mostly used for math but occasionally wraps body text too). We therefore
    accept any paragraph that has non-trivial prose content (heuristic: >= 25 chars
    of text, which excludes math-equation connectives like "or  or" or "From the
    formula sheet:") and isn't a heading.
    """
    doc = Document(docx_path)
    sections: dict[tuple[int, Optional[str]], list[str]] = {}
    current_key: Optional[tuple[int, Optional[str]]] = None
    PROSE_MIN_CHARS = 25
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "").strip()
        text = (p.text or "").strip()
        # Headings: detect by text-match (style is unreliable — the source docx applies
        # "VCAA body" to some headings like "Question 7d.").
        parsed = _parse_heading(text) if text else None
        if parsed is not None:
            current_key = parsed
            sections.setdefault(current_key, [])
            continue
        # Treat any other VCAA Heading-styled paragraph (e.g. "General comments",
        # "Specific information") as a section boundary that resets us out of any
        # question context.
        if style.startswith("VCAA Heading"):
            current_key = None
            continue
        if current_key is None:
            continue
        if not text or len(text) < PROSE_MIN_CHARS:
            continue
        sections[current_key].append(text)
    return {k: "\n\n".join(v).strip() for k, v in sections.items()}


# ─── Section A MC answers ─────────────────────────────────────────────

# Header row of the Section A answer-key table in VCAA examiner reports.
# Tolerant of minor whitespace/casing differences.
_MC_TABLE_HEADER_KEYS = ("Question", "Correct answer", "% A", "% B", "% C", "% D", "% E")


def _find_section_a_mc_table(docx_path: Path):
    """Locate the Section A answer-key table inside the examiner report docx.

    Returns the python-docx Table object, or None if not found. Identification is by
    header row content rather than position so it survives layout drift across years.
    """
    doc = Document(docx_path)
    for t in doc.tables:
        if not t.rows:
            continue
        header = [(c.text or "").strip() for c in t.rows[0].cells]
        if all(any(k.lower() in cell.lower() for cell in header) for k in _MC_TABLE_HEADER_KEYS):
            return t
    return None


def extract_mc_answers_from_docx(source_id: int, docx_path: Path,
                                  available_mc_qids: dict[int, str]) -> list[dict]:
    """Parse the Section A answer-key table; populate mc_correct + answers row.

    available_mc_qids maps question_number → qid for Section A questions. For 2023 P2
    that's {1: '2023-p2-sA-q1', ..., 20: '2023-p2-sA-q20'}.

    For each row we write:
      - questions.mc_correct (the single letter A-E)
      - answers.final_answer_md (compact markdown summary: correct letter + % distribution)
      - answers.commentary_md (the 'Comments' column text, often empty for routine questions)
      - answers.source_report_path (the .docx path, for traceability)
      - answers.answer_image_path = NULL (MC answers don't need a cropped image)
    Returns a list of per-row results.
    """
    table = _find_section_a_mc_table(docx_path)
    if table is None:
        return []

    header = [(c.text or "").strip() for c in table.rows[0].cells]
    # Map column names → index (tolerant of order drift).
    col = {h.lower(): i for i, h in enumerate(header)}
    q_col = col.get("question")
    correct_col = col.get("correct answer")
    comments_col = col.get("comments")
    pct_cols: dict[str, int] = {}
    for letter in ("A", "B", "C", "D", "E", "N/A"):
        idx = col.get(f"% {letter.lower()}")
        if idx is not None:
            pct_cols[letter] = idx
    if q_col is None or correct_col is None:
        return []

    rel_report = str(docx_path.relative_to(REPO_ROOT))
    results: list[dict] = []
    conn = connect()
    try:
        for row in table.rows[1:]:
            cells = [(c.text or "").strip() for c in row.cells]
            try:
                qn = int(cells[q_col])
            except ValueError:
                continue
            qid = available_mc_qids.get(qn)
            if qid is None:
                results.append({"qn": qn, "status": "no_question_row"})
                continue
            correct = cells[correct_col].strip().upper()
            if correct not in {"A", "B", "C", "D", "E"}:
                results.append({"qn": qn, "status": f"bad_correct={correct!r}"})
                continue
            comments = cells[comments_col].strip() if comments_col is not None else ""

            # Store a structured JSON blob in final_answer_md so the renderer can lay out
            # the correct answer + % distribution as styled chips rather than markdown
            # with a "←" arrow (which reads as a typo to students).
            distribution: list[dict] = []
            for letter, idx in pct_cols.items():
                v = cells[idx].strip()
                if not v:
                    continue
                distribution.append({"letter": letter, "pct": v})
            final_answer_md = json.dumps({
                "type": "mc_answer",
                "correct": correct,
                "distribution": distribution,
            })

            conn.execute(
                "update questions set mc_correct = ? where id = ?",
                (correct, qid),
            )
            conn.execute(
                """
                insert into answers (question_id, final_answer_md, commentary_md,
                                     source_report_path, answer_image_path)
                values (?, ?, ?, ?, NULL)
                on conflict (question_id) do update set
                  final_answer_md   = excluded.final_answer_md,
                  commentary_md     = excluded.commentary_md,
                  source_report_path = excluded.source_report_path,
                  answer_image_path = NULL
                """,
                (qid, final_answer_md, comments or "(no commentary)", rel_report),
            )
            results.append({"qn": qn, "qid": qid, "correct": correct,
                            "comments_len": len(comments), "status": "ok"})
        conn.commit()
    finally:
        conn.close()
    return results


def crop_mc_answer_table_rows(report_path: Path, source_id: int, *, force: bool = False) -> dict:
    """For an examiner-report PDF (or pre-converted docx-PDF), find the Section A
    answer-key table and crop each Q-row to a PNG so the human-in-the-loop audit
    UI can show the shaded cell directly.

    Writes to assets/mc_row_thumbs/<source_id>-q<NN>.png. Returns the map
    {question_number: relative_path}.
    """
    # If the report is docx, get the LibreOffice-rendered PDF.
    if report_path.suffix.lower() == ".docx":
        pdf_path = convert_docx_to_pdf(report_path)
    else:
        pdf_path = report_path

    doc = fitz.open(pdf_path)
    out: dict[int, str] = {}
    try:
        for pn in range(doc.page_count):
            pg = doc.load_page(pn)
            spans = []
            for blk in pg.get_text("dict").get("blocks", []):
                for ln in blk.get("lines", []):
                    for sp in ln.get("spans", []):
                        t = (sp.get("text") or "").strip()
                        if not t: continue
                        bb = sp["bbox"]
                        spans.append({"text": t, "x": (bb[0]+bb[2])/2, "y": (bb[1]+bb[3])/2,
                                      "x0": bb[0], "x1": bb[2], "y0": bb[1], "y1": bb[3]})
            # Detect this is an answer-table page by presence of "% A" / "% B" headers
            import re as _re
            letter_xs = []
            for s in spans:
                if _re.match(r"^%\s*[A-E]$", s["text"]):
                    letter_xs.append(s["x"])
            if len(letter_xs) < 3:
                continue

            leftmost_letter_x = min(letter_xs)
            # Q-number rows: digit spans 1..20 in the leftmost column of the table.
            q_rows: list[tuple[int, dict]] = []
            seen = set()
            for s in sorted(spans, key=lambda s: s["y"]):
                if not s["text"].isdigit():
                    continue
                qn = int(s["text"])
                if not (1 <= qn <= 20):
                    continue
                if s["x"] >= leftmost_letter_x - 15:
                    continue
                if qn in seen:
                    continue
                seen.add(qn)
                q_rows.append((qn, s))
            # Sort by y to get document order, then derive row y-bounds.
            q_rows.sort(key=lambda t: t[1]["y"])
            page_h = pg.rect.height
            for i, (qn, s) in enumerate(q_rows):
                y_top = s["y0"] - 3
                if i + 1 < len(q_rows):
                    y_bot = q_rows[i + 1][1]["y0"] - 2
                else:
                    # Last row on this page — bound by table bottom (look for next non-row content or page edge)
                    y_bot = min(s["y1"] + 60, page_h)
                # x-range: from the Q-column left edge to a bit past the rightmost %E column
                x_left = max(0, min(s["x0"] for _, s in q_rows) - 5)
                rightmost_letter_x = max(letter_xs)
                x_right = min(pg.rect.width, rightmost_letter_x + 60)
                out_path = MC_ROW_THUMBS_DIR / f"{source_id}-q{qn:02d}.png"
                if out_path.exists() and not force:
                    out[qn] = str(out_path.relative_to(REPO_ROOT))
                    continue
                clip = fitz.Rect(x_left, y_top, x_right, y_bot)
                pix = pg.get_pixmap(clip=clip, dpi=200)
                pix.save(out_path)
                out[qn] = str(out_path.relative_to(REPO_ROOT))
    finally:
        doc.close()
    return out


def _realign_to(target_header: list, src_table: list[list]) -> list[list]:
    """Re-map src_table's rows onto target_header's column count.

    PyMuPDF can extract the answer-key table with different sub-column counts on
    different pages (e.g. 2016 P2: page 2's E-column spans 1 cell, page 3's spans
    3 cells). We collapse each row to one value per NAMED header column — taking
    the first non-empty value within the source's span — then place each value at
    the same named column's position in the target. Within-span sub-positions are
    discarded; the percentage values don't carry sub-cell semantics anyway.
    """
    def named_positions(hdr):
        return [(i, (c or "").strip().lower()) for i, c in enumerate(hdr) if (c or "").strip()]

    tgt_named = named_positions(target_header)
    src_named = named_positions(src_table[0])
    if len(tgt_named) != len(src_named):
        return src_table[1:]  # last-resort: drop header, append as-is (may misalign)

    out_width = len(target_header)
    aligned: list[list] = []
    for row in src_table[1:]:
        new_row = [None] * out_width
        for k, (sidx, _) in enumerate(src_named):
            end = src_named[k + 1][0] if k + 1 < len(src_named) else len(row)
            value = None
            for i in range(sidx, min(end, len(row))):
                v = row[i]
                if v is not None and (str(v).strip() != ""):
                    value = v
                    break
            if value is not None:
                tgt_idx = tgt_named[k][0]
                new_row[tgt_idx] = value
        aligned.append(new_row)
    return aligned


def extract_mc_answers_from_pdf(source_id: int, pdf_path: Path,
                                  available_mc_qids: dict[int, str]) -> list[dict]:
    """PDF-native counterpart to `extract_mc_answers_from_docx`.

    2016–2019 examiner reports indicate the correct answer via cell shading on the
    percentage cell rather than via a labeled "Correct answer" column. We approximate
    by picking the option with the highest percentage. When the top two are close
    (margin < 10 points), we still pick the higher but log a review_queue row so the
    user can audit by hand.
    """
    src = fitz.open(pdf_path)
    try:
        # The answer-key table can span multiple pages. Some years (e.g. 2019) split
        # the header across 2 rows because of wrapping in the "% No answer" cell —
        # row 0 is empty, row 1 has the actual labels. So we scan the first ~3 rows
        # of each candidate table for a row containing "% A", "% B", ..., and treat
        # that row as the header (discarding rows above it).
        def _find_header_row(data: list[list]) -> Optional[int]:
            for ri in range(min(3, len(data))):
                cells = [(c or "").strip().lower() for c in data[ri]]
                joined = " ".join(cells)
                if all(f"% {L.lower()}" in joined for L in ("A", "B", "C", "D")):
                    return ri
            return None

        chosen_table = None
        for pn in range(src.page_count):
            pg = src.load_page(pn)
            try:
                tables = pg.find_tables()
            except Exception:
                continue
            for tbl in tables:
                data = tbl.extract()
                if not data:
                    continue
                hdr_idx = _find_header_row(data)
                if chosen_table is None:
                    if hdr_idx is not None:
                        chosen_table = data[hdr_idx:]
                elif hdr_idx is not None:
                    # Continuation table with repeated header — strip rows up to and
                    # including the header, then re-align to chosen_table's columns.
                    chosen_table.extend(_realign_to(chosen_table[0], data[hdr_idx:]))
        if chosen_table is None:
            return [{"status": "no_mc_table_found"}]
    finally:
        src.close()

    # PyMuPDF preserves the table's visual cell structure. Each "% A" header in
    # 2016–2019 reports spans 3 sub-columns, with the actual percentage value
    # appearing in any one of them (depending on the cell's horizontal centring).
    # We record the header column AND its span so we can scan across sub-cells.
    header = [(c or "").strip().lower() for c in chosen_table[0]]
    pct_spans: dict[str, tuple[int, int]] = {}   # letter -> (start_idx, end_idx_exclusive)
    q_col: Optional[int] = None
    comments_col: Optional[int] = None
    # Walk the header recording each non-empty cell and its run-length to the next
    # non-empty header. None/empty cells are merged into the preceding header.
    named_cols: list[tuple[int, str]] = [(i, h) for i, h in enumerate(header) if h]
    for k, (idx, h) in enumerate(named_cols):
        next_idx = named_cols[k + 1][0] if k + 1 < len(named_cols) else len(header)
        if h in ("question", "qu", "q"):
            q_col = idx
        elif h == "comments":
            comments_col = idx
        elif h.startswith("% "):
            letter = h[2:].strip().upper()
            if letter and letter[0] in "ABCDEN":
                key = letter[0] if letter[0] != "N" else "N/A"
                pct_spans[key] = (idx, next_idx)
    if q_col is None or not pct_spans:
        return [{"status": "header_unrecognised", "header": header}]

    # Some Q-rows split across 2 PyMuPDF rows when the Comments cell has multi-line
    # content (e.g. inline math expressions across several lines). Merge any row
    # whose Question column is empty into the prior Q-numbered row.
    merged_rows: list[list] = []
    for row in chosen_table[1:]:
        q_cell = (row[q_col] if q_col < len(row) else "") or ""
        if q_cell.strip().isdigit():
            merged_rows.append([(c if c is not None else "") for c in row])
        elif merged_rows:
            for i, c in enumerate(row):
                if c and i < len(merged_rows[-1]) and not merged_rows[-1][i]:
                    merged_rows[-1][i] = c

    rel_report = str(pdf_path.relative_to(REPO_ROOT))
    results: list[dict] = []
    conn = connect()
    try:
        for row in merged_rows:
            cells = [(c or "").strip() for c in row]
            try:
                qn = int(cells[q_col])
            except (ValueError, IndexError):
                continue
            qid = available_mc_qids.get(qn)
            if qid is None:
                results.append({"qn": qn, "status": "no_question_row"})
                continue

            # Highest-percentage heuristic for the correct answer. For each letter,
            # scan the header's sub-column span and take the first non-empty value.
            def _pct_in_span(start: int, end: int) -> Optional[str]:
                for i in range(start, min(end, len(cells))):
                    v = cells[i].strip()
                    if v:
                        return v
                return None

            distribution: list[dict] = []
            pcts: list[tuple[str, float]] = []
            for letter, (s, e) in pct_spans.items():
                v = _pct_in_span(s, e)
                if v is None:
                    continue
                distribution.append({"letter": letter, "pct": v})
                if letter == "N/A":
                    continue
                try:
                    pcts.append((letter, float(v)))
                except ValueError:
                    continue
            if not pcts:
                results.append({"qn": qn, "status": "no_percentages"})
                continue
            pcts.sort(key=lambda lp: lp[1], reverse=True)
            correct = pcts[0][0]
            margin = pcts[0][1] - (pcts[1][1] if len(pcts) >= 2 else 0)
            if margin < 10:
                conn.execute(
                    "insert into review_queue (question_id, source_id, page, reason, detail) values (?, ?, ?, ?, ?)",
                    (qid, source_id, None, "mc_correct_low_margin",
                     f"Q{qn} top pick is {correct} ({pcts[0][1]}%) but runner-up is {pcts[1][0]} ({pcts[1][1]}%) — margin {margin:.1f}, verify by inspecting the cell shading in {rel_report}."),
                )

            comments = cells[comments_col].strip() if (comments_col is not None and comments_col < len(cells)) else ""
            final_answer_md = json.dumps({
                "type": "mc_answer",
                "correct": correct,
                "distribution": distribution,
            })
            conn.execute(
                "update questions set mc_correct = ? where id = ?",
                (correct, qid),
            )
            conn.execute(
                """
                insert into answers (question_id, final_answer_md, commentary_md,
                                     source_report_path, answer_image_path)
                values (?, ?, ?, ?, NULL)
                on conflict (question_id) do update set
                  final_answer_md   = excluded.final_answer_md,
                  commentary_md     = excluded.commentary_md,
                  source_report_path = excluded.source_report_path,
                  answer_image_path = NULL
                """,
                (qid, final_answer_md, comments or "(no commentary)", rel_report),
            )
            results.append({"qn": qn, "qid": qid, "correct": correct,
                            "margin": margin, "status": "ok"})
        conn.commit()
    finally:
        conn.close()
    return results


# ─── question-id resolution ───────────────────────────────────────────

def resolve_question_id(source_id: int, question_number: int,
                        part: Optional[str], available_qids: dict) -> Optional[str]:
    """Map (question_number, part) → existing questions.id. Uses the leaf-row convention.

    For a heading 'Question 1' that has sub-parts a, b in our DB, the heading is
    "the whole question" — we attach the report section to the STEM row
    (part=null) if it exists, or to the only leaf if there are no parts.
    For 'Question 1a.', we attach to the 1.a leaf row.

    Examiner-report quirk (2020): the report may subdivide a paper-part into
    answer-step labels like '1b.i', '1b.ii' even though the actual paper has Q1.b
    as a single part. When (qn, 'b.i') doesn't match, progressively strip the
    trailing dot-component ('b.i' → 'b') until a match is found.
    """
    key = (question_number, part)
    if key in available_qids:
        return available_qids[key]
    # Progressively strip trailing dot-components: 'b.i.iv' → 'b.i' → 'b'.
    if part and "." in part:
        stripped = part
        while "." in stripped:
            stripped = stripped.rsplit(".", 1)[0]
            if (question_number, stripped) in available_qids:
                return available_qids[(question_number, stripped)]
    # If heading is "Question N" with no part, try a stem row, else the unique whole.
    if part is None:
        stem_id = available_qids.get((question_number, None))
        if stem_id:
            return stem_id
    return None


# ─── orchestrator ─────────────────────────────────────────────────────

def extract_paper_answers(year: int, paper: int, *, force: bool = False) -> dict:
    conn = connect()
    try:
        src = conn.execute(
            "select id, report_path from sources where year = ? and paper = ?",
            (year, paper),
        ).fetchone()
        if src is None:
            raise RuntimeError(f"no source for {year} paper {paper}")
        source_id = src["id"]
        if not src["report_path"]:
            raise RuntimeError(f"no examiner report attached for {year} paper {paper}")
        report_path = REPO_ROOT / src["report_path"]
        if report_path.suffix.lower() not in (".docx", ".pdf"):
            raise RuntimeError(
                f"Unsupported examiner-report format: {report_path.suffix}"
            )
        is_pdf = report_path.suffix.lower() == ".pdf"
        # Collect questions for this source, keyed by (qnum, part). Section A MCs
        # get a separate map indexed by question_number alone since the answer-key
        # table is one row per question (no parts).
        qrows = conn.execute(
            "select id, question_number, part, section from questions where source_id = ?",
            (source_id,),
        ).fetchall()
        available_qids: dict[tuple[int, Optional[str]], str] = {}
        available_mc_qids: dict[int, str] = {}
        for r in qrows:
            if r["section"] == "A":
                available_mc_qids[r["question_number"]] = r["id"]
            else:
                available_qids[(r["question_number"], r["part"])] = r["id"]
    finally:
        conn.close()

    if is_pdf:
        pdf_path = report_path
        commentary_map = pdf_commentary_by_question(pdf_path)
    else:
        pdf_path = convert_docx_to_pdf(report_path, force=force)
        commentary_map = docx_commentary_by_question(report_path)
    headings = find_headings(pdf_path)
    sections = section_ranges(pdf_path, headings)

    results = []
    skipped = []
    # First pass: resolve each docx section to a qid (with fallback for the
    # 2020-style report subdivision where 'Q1b.i' / 'Q1b.ii' both belong to Q1.b).
    resolved: list[tuple[str, dict, list, tuple]] = []  # (qid, heading, ranges, original_key)
    for heading, ranges in sections:
        qn = heading["question_number"]
        part = heading["part"]
        qid = resolve_question_id(0, qn, part, available_qids)
        if qid is None:
            skipped.append({"heading": heading["text"], "reason": "no matching question_id",
                            "qn": qn, "part": part})
            continue
        resolved.append((qid, heading, list(ranges), (qn, part)))

    # Second pass: merge consecutive sections that resolved to the same qid.
    # This handles the report-subdivides-a-single-part case: Q1.b.i + Q1.b.ii both
    # resolve to '...-q1-b' and we want ONE combined image + concatenated commentary,
    # not the second one clobbering the first.
    merged: list[dict] = []
    for qid, heading, ranges, orig_key in resolved:
        if merged and merged[-1]["qid"] == qid:
            merged[-1]["ranges"].extend(ranges)
            merged[-1]["headings"].append(heading["text"])
            merged[-1]["orig_keys"].append(orig_key)
        else:
            merged.append({"qid": qid, "headings": [heading["text"]],
                           "ranges": ranges, "orig_keys": [orig_key]})

    for entry in merged:
        qid = entry["qid"]
        ranges = entry["ranges"]
        out_image = ANSWERS_DIR / f"{qid}.png"
        crop_section(pdf_path, ranges, out_image)

        # Concat commentary across all the original (qn, part) keys that mapped here.
        chunks = []
        for k in entry["orig_keys"]:
            txt = commentary_map.get(k, "").strip()
            if txt:
                chunks.append(txt)
        commentary = "\n\n".join(chunks)

        rel_image = str(out_image.relative_to(REPO_ROOT))
        rel_report = str(report_path.relative_to(REPO_ROOT))

        conn = connect()
        try:
            conn.execute(
                """
                insert into answers (question_id, final_answer_md, commentary_md,
                                     source_report_path, answer_image_path)
                values (?, ?, ?, ?, ?)
                on conflict (question_id) do update set
                  commentary_md      = excluded.commentary_md,
                  source_report_path = excluded.source_report_path,
                  answer_image_path  = excluded.answer_image_path
                """,
                (qid, None, commentary or "(no body text)", rel_report, rel_image),
            )
            conn.commit()
        finally:
            conn.close()

        results.append({
            "qid": qid,
            "heading": "; ".join(entry["headings"]),
            "image": rel_image,
            "commentary_chars": len(commentary),
            "pages_spanned": len(ranges),
        })

    # Section A MC answers: pulled from a single table in the report, no PDF cropping.
    mc_results: list[dict] = []
    if available_mc_qids:
        if is_pdf:
            mc_results = extract_mc_answers_from_pdf(source_id, report_path, available_mc_qids)
        else:
            mc_results = extract_mc_answers_from_docx(source_id, report_path, available_mc_qids)

    return {
        "year": year,
        "paper": paper,
        "headings_found": len(headings),
        "answers_inserted": len(results),
        "mc_answers_inserted": sum(1 for r in mc_results if r.get("status") == "ok"),
        "skipped": skipped,
        "results": results,
        "mc_results": mc_results,
    }


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--year", type=int, required=True)
    ap.add_argument("--paper", type=int, required=True, choices=[1, 2])
    ap.add_argument("--force", action="store_true",
                    help="re-convert docx→pdf even if cached")
    args = ap.parse_args()
    print(json.dumps(extract_paper_answers(args.year, args.paper, force=args.force), indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
